import docx
from utils_functions.create_title import create_title
from utils_functions.create_table import create_table, add_table_row
from utils_functions.create_image import create_image


def create_word_template_with_image(output_path, img_paths):
    doc = docx.Document()

    create_title(doc, "足踝检测评估报告", 1)
    # Add the contents for the first row of the table
    table = create_table(doc, 0, 3)
    add_table_row(table, ["姓名：\t", "年龄：\t", "性别："])

    create_image(doc, img_paths[0], 6)
    doc.add_page_break()

    # Add the title on the second page
    create_title(doc, "实际足部受力成像情况：", 2)
    create_image(doc, img_paths[1], 6)  

    # Add a table to place two images horizontally
    compare_table = create_table(doc, 1, 2)
    add_image_with_title(doc, img_paths[2], "正常足底受力", compare_table.cell(0, 0))
    add_image_with_title(doc, img_paths[3], "实际足底受力", compare_table.cell(0, 1))

    doc.add_page_break() 

    create_title(doc, "足跟内外翻情况对比：", 2)
    compare_table2 = create_table(doc, 1, 2)
    add_image_with_title(doc, img_paths[4], "正常后跟内外翻情况", compare_table2.cell(0, 0))
    add_image_with_title(doc, img_paths[5], "实际后跟内外翻情况", compare_table2.cell(0, 1))

    # Add text after the tables
    create_title(doc, "初步诊断：", 2)
    doc.add_paragraph("足弓低平且足后跟外偏，提示可能有运动损伤风险。处于生长发育阶段的儿童及青少年，应使用有矫正效果的鞋垫，并搭配合适的鞋类，日常应加强足部及腿部肌肉锻炼、拉伸及放松，以改善足弓扁平及足部外翻的情况，防止足部问题因体重增长及不良步态导致进行性加重，从而影响膝、骨盆与脊柱的正常生物力线；\n\n足部问题常伴有体态姿势不良等问题的出现，处于生发育阶段的儿童及青少年应每3个月到半年进行一次全面的脊柱、体态、骨龄及足部的相关检查。")
    create_image(doc, img_paths[6], 6)

    doc.save(output_path)


def add_image_with_title(doc, img_path, title_text, cell):
    cell.width = docx.shared.Inches(3)
    cell_img = cell.add_paragraph().add_run()
    create_image(cell_img, img_path, 3)

    title = cell.add_paragraph()
    title.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(title_text)

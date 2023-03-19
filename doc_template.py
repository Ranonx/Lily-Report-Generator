import docx
from utils_functions.create_image import create_image

def add_image_with_title(doc, img_path, title_text, cell):
    cell.width = docx.shared.Inches(3)
    cell_img = cell.add_paragraph().add_run()
    create_image(cell_img, img_path, 3)

    title = cell.add_paragraph()
    title.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(title_text)
